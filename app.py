import os
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def bold_keywords_in_doc(source_path, destination_path, keywords):
    doc = Document(source_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Preserve the original run formatting
            run_style = run.style
            run_font = run.font
            original_text = run.text
            run.text = ""

            words = original_text.split(' ')
            for word in words:
                # Check if the word is a keyword, ignoring case and punctuation
                stripped_word = word.strip('.,;?!:"\'()').lower()
                if any(stripped_word == keyword.lower() for keyword in keywords):
                    # If it's a keyword, add it with bold formatting
                    new_run = paragraph.add_run(word + ' ')
                    new_run.bold = True
                    new_run.style = run_style
                    new_run.font.name = run_font.name
                    new_run.font.size = run_font.size
                    new_run.font.color.rgb = run_font.color.rgb
                else:
                    # If it's not a keyword, add it with original formatting
                    new_run = paragraph.add_run(word + ' ')
                    new_run.style = run_style
                    new_run.font.name = run_font.name
                    new_run.font.size = run_font.size
                    new_run.font.color.rgb = run_font.color.rgb

    doc.save(destination_path)



# Comprehensive list of keywords to be bolded in the document
keywords = [
    "Python", "Django", "Software Engineer", "web applications", "scalable", "maintainable",
    "collaborating", "cross-functional teams", "high-quality software solutions", "technical specifications",
    "clean", "efficient", "well-documented code", "code reviews", "code quality", "maintainability",
    "scalability", "application performance", "automated testing", "continuous integration",
    "robust", "fast-paced environment", "design", "develop", "best practices", "coding standards",
    "latest trends and technologies", "QA engineers", "HTML", "CSS", "JavaScript", "React", "Angular",
    "PostgreSQL", "MySQL", "MongoDB", "Git", "problem-solving", "communication skills", "teamwork",
    "Agile development", "AWS", "Azure", "Google Cloud Platform", "Docker", "Kubernetes", 
    "open-source projects", "developer community"
]

# Paths to the source and destination documents from environment variables
source_path = os.getenv('SOURCE_DOC')
destination_path = os.getenv('DESTINATION_DOC')

# Function call to bold the keywords
bold_keywords_in_doc(source_path, destination_path, keywords)




# import os
# from docx import Document
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()

# def bold_keywords_in_doc(source_path, destination_path, keywords):
#     doc = Document(source_path)
    
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             # Splitting the text in the run by spaces to check each word against the keywords
#             words = run.text.split()
#             for i, word in enumerate(words):
#                 # Checking each word against the keywords list
#                 if any(word.lower().strip('.,!?;:') == keyword.lower() for keyword in keywords):
#                     # If the word is a keyword, we bold it
#                     words[i] = '**' + word + '**'
#             # Reconstruct the run's text while preserving original formatting
#             run.text = ' '.join(words)

#     doc.save(destination_path)

# # Comprehensive list of keywords to be bolded in the document
# keywords = [
#     "Python", "Django", "Software Engineer", "web applications", "scalable", "maintainable",
#     "collaborating", "cross-functional teams", "high-quality software solutions", "technical specifications",
#     "clean", "efficient", "well-documented code", "code reviews", "code quality", "maintainability",
#     "scalability", "application performance", "automated testing", "continuous integration",
#     "robust", "fast-paced environment", "design", "develop", "best practices", "coding standards",
#     "latest trends and technologies", "QA engineers", "HTML", "CSS", "JavaScript", "React", "Angular",
#     "PostgreSQL", "MySQL", "MongoDB", "Git", "problem-solving", "communication skills", "teamwork",
#     "Agile development", "AWS", "Azure", "Google Cloud Platform", "Docker", "Kubernetes", 
#     "open-source projects", "developer community"
# ]

# # Paths to the source and destination documents from environment variables
# source_path = os.getenv('SOURCE_DOC')
# destination_path = os.getenv('DESTINATION_DOC')

# # Function call to bold the keywords
# bold_keywords_in_doc(source_path, destination_path, keywords)




# import os
# from docx import Document
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()

# def bold_keywords_in_doc(source_path, destination_path, keywords):
#     # Load the Word document from the specified path
#     doc = Document(source_path)

#     # Iterate through each paragraph and run in the document
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             # Check if any keyword exists in the run text, ignoring case
#             if any(keyword.lower() in run.text.lower() for keyword in keywords):
#                 # If keyword found, set the text to bold
#                 run.bold = True

#     # Save the modified document to the destination path
#     doc.save(destination_path)

# # Comprehensive list of keywords to be bolded in the document
# keywords = [
#     "Python", "Django", "Software Engineer", "web applications", "scalable", "maintainable",
#     "collaborating", "cross-functional teams", "high-quality software solutions", "technical specifications",
#     "clean", "efficient", "well-documented code", "code reviews", "code quality", "maintainability",
#     "scalability", "application performance", "automated testing", "continuous integration",
#     "robust", "fast-paced environment", "design", "develop", "best practices", "coding standards",
#     "latest trends and technologies", "QA engineers", "HTML", "CSS", "JavaScript", "React", "Angular",
#     "PostgreSQL", "MySQL", "MongoDB", "Git", "problem-solving", "communication skills", "teamwork",
#     "Agile development", "AWS", "Azure", "Google Cloud Platform", "Docker", "Kubernetes", 
#     "open-source projects", "developer community"
# ]

# # Paths to the source and destination documents from environment variables
# source_path = os.getenv('SOURCE_DOC')
# destination_path = os.getenv('DESTINATION_DOC')

# # Function call to bold the keywords
# bold_keywords_in_doc(source_path, destination_path, keywords)







# import os
# from docx import Document
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()

# def bold_keywords_in_doc(source_path, destination_path, keywords):
#     # Load the Word document from the specified path
#     doc = Document(source_path)

#     # Iterate through each paragraph and run in the document
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             # Check if any keyword exists in the run text, ignoring case
#             if any(keyword.lower() in run.text.lower() for keyword in keywords):
#                 # If keyword found, set the text to bold
#                 run.bold = True

#     # Save the modified document to the destination path
#     doc.save(destination_path)

# # List of keywords to be bolded in the document
# keywords = [
#     "Python", "Django", "Software Engineer", "web applications", "scalable", "maintainable",
#     "collaborating", "cross-functional teams", "high-quality software solutions", "technical specifications",
#     "clean", "efficient", "well-documented code", "code reviews", "code quality", "maintainability",
#     "scalability", "application performance", "automated testing", "continuous integration"
# ]

# # Paths to the source and destination documents from environment variables
# source_path = os.getenv('SOURCE_DOC')
# destination_path = os.getenv('DESTINATION_DOC')

# # Function call to bold the keywords
# bold_keywords_in_doc(source_path, destination_path, keywords)
